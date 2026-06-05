from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont, ImageStat


WORKDIR = Path(r"C:\Users\Chels\Documents\New project")
TEMPLATE = Path(r"D:\Downloads\DV2 - Assignment Template v3.0.docx")
OUTPUT = WORKDIR / "DV2 - Assignment Template v3.0 - filled.docx"
ASSET_DIR = WORKDIR / "submission_assets"

SKETCH_PLACEHOLDER = WORKDIR / "sketch-placeholder.png"
SKETCH_IMAGE = Path(
    r"C:\Users\Chels\Documents\WeChat Files\wxid_qh0jootfa7jk22\FileStorage\Temp\66094347a3d78a9012432a1ecd3eca6.jpg"
)
SKETCH_COPY = ASSET_DIR / "great-barrier-reef-sketch.jpg"
SKETCH_PDF = ASSET_DIR / "great-barrier-reef-sketch.pdf"
SKETCH_GITHUB_URL = (
    "https://raw.githubusercontent.com/Viatin11116/New-project/main/"
    "submission_assets/great-barrier-reef-sketch.pdf"
)
SITE_CAPTURE = WORKDIR / "site-full-edge-long.png"
SITE_FULL = WORKDIR / "site-full.png"
SITE_CROPPED = WORKDIR / "site-full-cropped.png"
SITE_FULL_FIT = WORKDIR / "site-full-single.png"


def get_font(size: int):
    candidates = [
        Path(r"C:\Windows\Fonts\arial.ttf"),
        Path(r"C:\Windows\Fonts\segoeui.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def make_sketch_placeholder():
    image = Image.new("RGB", (1400, 900), "white")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((30, 30, 1370, 870), radius=24, outline="#2E6F73", width=4)

    title_font = get_font(48)
    body_font = get_font(28)

    title = "Sketch Screenshot Placeholder"
    body = (
        "Sketch PDF found in Downloads:\n"
        "A_wireframe_digital_sketch_displays_wireframes_for.pdf\n\n"
        "This template is filled, but the PDF preview could not be rendered\n"
        "automatically in headless mode.\n\n"
        "If needed, replace this placeholder with an exported screenshot\n"
        "from your sketch PDF before final submission."
    )

    draw.text((80, 110), title, fill="#12343B", font=title_font)
    draw.multiline_text((80, 230), body, fill="#3C4F52", font=body_font, spacing=16)
    image.save(SKETCH_PLACEHOLDER)


def prepare_sketch_assets():
    ASSET_DIR.mkdir(exist_ok=True)

    if SKETCH_IMAGE.exists():
        sketch = Image.open(SKETCH_IMAGE).convert("RGB")
        sketch.save(SKETCH_COPY, format="JPEG", quality=95)
        sketch.save(SKETCH_PDF, "PDF", resolution=200.0)
        return SKETCH_IMAGE

    make_sketch_placeholder()
    placeholder = Image.open(SKETCH_PLACEHOLDER).convert("RGB")
    placeholder.save(SKETCH_COPY, format="JPEG", quality=95)
    placeholder.save(SKETCH_PDF, "PDF", resolution=200.0)
    return SKETCH_PLACEHOLDER


def make_site_full_fit():
    source_path = SITE_CAPTURE if SITE_CAPTURE.exists() else SITE_FULL
    source = Image.open(source_path).convert("RGB")
    gray = source.convert("L")
    width, height = gray.size
    last_content_row = height - 1

    for y in range(height - 1, -1, -1):
        row = gray.crop((120, y, width - 120, y + 1))
        stat = ImageStat.Stat(row)
        if stat.stddev[0] > 2:
            last_content_row = y
            break

    cropped_height = min(height, last_content_row + 40)
    source = source.crop((0, 0, width, cropped_height))
    source.save(SITE_CROPPED)

    canvas_w, canvas_h = 1400, 2000
    margin = 36

    available_w = canvas_w - margin * 2
    available_h = canvas_h - margin * 2
    scale = min(available_w / source.width, available_h / source.height)

    new_w = max(1, int(source.width * scale))
    new_h = max(1, int(source.height * scale))
    resized = source.resize((new_w, new_h))

    canvas = Image.new("RGB", (canvas_w, canvas_h), "white")
    x = (canvas_w - new_w) // 2
    y = (canvas_h - new_h) // 2
    canvas.paste(resized, (x, y))
    canvas.save(SITE_FULL_FIT)


def insert_image(paragraph, image_path: Path, width_inches: float):
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def main():
    sketch_to_use = prepare_sketch_assets()
    make_site_full_fit()

    doc = Document(str(TEMPLATE))

    # Header table.
    table = doc.tables[0]
    table.rows[0].cells[1].text = "Jiaxu Dong"
    table.rows[1].cells[1].text = "35823194"
    table.rows[2].cells[1].text = "17"

    # Part A
    doc.paragraphs[6].text = "https://viatin11116.github.io/New-project/"

    # Part B
    doc.paragraphs[14].text = (
        "This project is in environmental monitoring and coastal water-quality "
        "management. It examines pollution pressure in Great Barrier Reef inshore "
        "waters, especially sediment, nutrients and pesticide-related runoff that "
        "reaches reef-connected coastal water."
    )
    doc.paragraphs[19].text = (
        "I chose the topic because Reef pollution is often discussed in broad terms, "
        "but general readers rarely see how it varies by place, pollutant and time. "
        "The target audience is a public audience rather than technical specialists, "
        "so the page is written as an explanatory story instead of a technical dashboard. "
        "Readers should be able to identify where hotspots occur, which pollutants matter, "
        "when runoff signals strengthen, what ecological responses appear and whether "
        "management targets are being met."
    )
    doc.paragraphs[24].text = (
        "The visualisation uses only official data. The main sources are the AIMS "
        "inshore water-quality monitoring database, AIMS Marine Monitoring Program "
        "hard coral cover records, and Reef Plan report card and target reporting for "
        "2021-2022. I selected these sources because they connect short-term "
        "water-quality observations, longer-term ecological response and formal "
        "management assessment. I cleaned and reorganised the data into analysis-ready "
        "tables for Vega-Lite by filtering to relevant inshore records, calculating site "
        "or regional medians, grouping map values into percentile classes and converting "
        "reported current-versus-target reductions into target-completion percentages."
    )
    doc.paragraphs[29].text = (
        "I used different idioms for different tasks. A map answers the first question of "
        "where suspended-solids hotspots appear. Small-multiple bar charts compare regional "
        "rankings across DIN, suspended solids and chlorophyll a, showing that hotspots "
        "change by indicator. A heatmap summarises Reef Plan pollutant grades by region and "
        "category. Line charts show monthly and longer-term temporal patterns, while "
        "scatterplots show relationships between water-quality measures and ecological "
        "response variables. Progress bars communicate management completion rates quickly "
        "for non-specialist readers. The page is structured as a chapter-based story so "
        "readers move from monitoring scope to pollutant pattern, time signal, ecological "
        "consequence and management progress. Each chart is a separate Vega-Lite JSON "
        "specification loaded through a custom JavaScript embedding script, which made the "
        "project modular and easy to revise."
    )

    # Part C
    doc.paragraphs[30].style = "Heading 3"
    doc.paragraphs[30].text = "AI use declaration"
    doc.paragraphs[31].text = (
        "Generative AI tools were used to support code iteration, layout revision, and "
        "wording refinement. All topic selection, source checking, data-use decisions, "
        "chart interpretation, and final submitted content were decided by the student."
    )

    doc.paragraphs[36].text = SKETCH_GITHUB_URL
    insert_image(doc.paragraphs[41], sketch_to_use, 6.0)

    # Part D
    insert_image(doc.paragraphs[45], SITE_FULL_FIT, 6.3)

    # Light formatting cleanup for filled answers.
    for idx in [6, 14, 19, 24, 29, 31, 36]:
        for run in doc.paragraphs[idx].runs:
            run.font.size = Pt(10.5)

    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
